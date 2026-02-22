"""
Karoo v2.0 Exporter
Generates TXT, DOCX, and PDF reports. Enhanced formatting.
"""
import io
from datetime import datetime
from typing import Dict, Any, Optional


def export_to_txt(results: Dict[str, Any]) -> str:
    summary = results.get("summary", {})
    score = summary.get("overall_score", 0)
    sep = "=" * 68
    thin = "-" * 68

    icons = {
        "algorithm_breaker": "ðŸŽ¯", "sa_specialist": "ðŸ‡¿ðŸ‡¦", "global_setter": "ðŸŒ",
        "recruiter_scanner": "ðŸ‘ï¸ ", "hiring_manager": "ðŸ’¼", "semantic_matcher": "ðŸ“Š",
        "compliance_guardian": "âš–ï¸ ", "future_architect": "ðŸš€",
        "interview_coach": "ðŸŽ¤", "salary_intelligence": "ðŸ’°",
    }

    lines = [
        sep,
        "  Karoo v2.0 â€” OPTIMIZATION REPORT",
        f"  Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
        f"  AI Provider: {results.get('llm_provider','Rule-Based')} â€” {results.get('llm_model','N/A')}",
        f"  AI-Powered Agents: {summary.get('ai_powered_count', 0)}/10",
        sep, "",
        f"  OVERALL SCORE:           {score}/100",
        f"  INTERVIEW PROBABILITY:   {summary.get('interview_probability',0)}%",
        f"  RECOMMENDED VARIANT:     {summary.get('recommended_variant','BALANCED')}",
        f"  VERDICT:                 {summary.get('verdict','')}",
        f"  STRONGEST AREA:          {summary.get('strongest_area','').replace('_',' ').title()}",
        f"  WEAKEST AREA:            {summary.get('weakest_area','').replace('_',' ').title()}",
        "", sep, "  AGENT SCORES", sep,
    ]

    for name, s in summary.get('agent_scores', {}).items():
        rating = "EXCELLENT âœ“âœ“" if s>=85 else "STRONG âœ“" if s>=75 else "ADEQUATE ~" if s>=60 else "NEEDS WORK !"
        bar = "â–ˆ"*(s//10) + "â–‘"*(10-s//10)
        icon = icons.get(name, "â€¢ ")
        lines.append(f"  {icon} {name.replace('_',' ').title():<32} {s:>3}/100  [{bar}]  {rating}")

    lines += ["", sep, "  PRIORITY ACTION ITEMS", sep, ""]
    for i, item in enumerate(results.get('action_items', [])[:18], 1):
        priority = "ðŸ”´ CRITICAL" if i<=4 else "ðŸŸ¡ IMPORTANT" if i<=9 else "ðŸŸ¢ HELPFUL"
        lines.append(f"  {i:>2}. [{priority}] {item}")

    for vk, vl in [("balanced","BALANCED â­"),("ats_max","ATS-MAX"),("creative","CREATIVE")]:
        content = results.get('cv_variants', {}).get(vk, '')
        if content:
            lines += ["", sep, f"  CV VARIANT: {vl}", sep, "", content]

    cl = results.get('cover_letter', '')
    if cl:
        lines += ["", sep, "  COVER LETTER", sep, "", cl]

    lines += ["", sep, "  DETAILED AGENT REPORTS", sep]
    for name, data in results.get('agent_results', {}).items():
        lines += ["", f"  {thin}", f"  {name.replace('_',' ').upper()} â€” {data.get('score',0)}/100", f"  {thin}"]
        lines.append("  Findings:")
        for f in data.get('findings', []): lines.append(f"    â€¢ {f}")
        lines.append("  Recommendations:")
        for r in data.get('recommendations', []): lines.append(f"    â†’ {r}")
        if data.get('optimized_content'):
            lines += ["  AI-Generated Improvement:", f"  {data['optimized_content'][:500]}"]

    return '\n'.join(lines)


def export_to_docx(results: Dict[str, Any], variant: str = "balanced") -> Optional[bytes]:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        summary = results.get("summary", {})
        score = summary.get("overall_score", 0)

        # Title
        title = doc.add_heading('Karoo v2.0 â€” Optimization Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
        doc.add_paragraph(f"AI Provider: {results.get('llm_provider','Rule-Based')} ({results.get('llm_model','N/A')})")

        # Score summary table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        cells = [
            ("Overall ATS Score", f"{score}/100"),
            ("Interview Probability", f"{summary.get('interview_probability',0)}%"),
            ("Recommended Variant", summary.get('recommended_variant','BALANCED')),
        ]
        for i, (label, val) in enumerate(cells):
            if i < 3:
                r = table.rows[i]
                r.cells[0].text = label
                r.cells[0].paragraphs[0].runs[0].bold = True
                r.cells[1].text = val

        doc.add_paragraph()
        doc.add_paragraph(f"Verdict: {summary.get('verdict','')}")

        # Agent scores
        doc.add_heading('Agent Scores', level=1)
        for agent, s in summary.get('agent_scores', {}).items():
            p = doc.add_paragraph()
            label = agent.replace('_', ' ').title()
            bar = "â–ˆ" * (s // 10) + "â–‘" * (10 - s // 10)
            r = p.add_run(f"{label:<35} {s:>3}/100  [{bar}]")
            if s >= 80: r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            elif s >= 60: r.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
            else: r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
            r.font.name = 'Courier New'

        # Action items
        doc.add_heading('Priority Action Items', level=1)
        for i, item in enumerate(results.get('action_items', [])[:18], 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(item)

        # CV Variant
        variants = results.get('cv_variants', {})
        cv_content = variants.get(variant, variants.get('balanced', ''))
        if cv_content:
            doc.add_page_break()
            doc.add_heading(f'CV Variant: {variant.upper().replace("_","-")}', level=1)
            for line in cv_content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    if line.startswith('â•') or line.startswith('â•”') or line.startswith('â•‘'):
                        p.runs[0].bold = True
                        p.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        # Cover letter
        cl = results.get('cover_letter', '')
        if cl:
            doc.add_page_break()
            doc.add_heading('Cover Letter', level=1)
            for para in cl.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Interview Q&A
        interview = results.get('agent_results', {}).get('interview_coach', {})
        if interview.get('optimized_content'):
            doc.add_page_break()
            doc.add_heading('Interview Preparation', level=1)
            doc.add_paragraph(interview['optimized_content'][:3000])

        # Salary Intelligence
        salary = results.get('agent_results', {}).get('salary_intelligence', {})
        if salary.get('findings'):
            doc.add_heading('Salary Intelligence', level=1)
            for f in salary.get('findings', []):
                doc.add_paragraph(f"â€¢ {f}")
            if salary.get('optimized_content'):
                doc.add_paragraph("Negotiation Script:")
                doc.add_paragraph(salary['optimized_content'])

        # Full agent reports
        doc.add_page_break()
        doc.add_heading('Detailed Agent Reports', level=1)
        for name, data in results.get('agent_results', {}).items():
            doc.add_heading(f"{name.replace('_',' ').title()} â€” {data.get('score',0)}/100", level=2)
            for f in data.get('findings', []):
                doc.add_paragraph(f"â€¢ {f}")
            doc.add_paragraph("Recommendations:")
            for r_item in data.get('recommendations', []):
                doc.add_paragraph(f"â†’ {r_item}")
            if data.get('optimized_content'):
                doc.add_paragraph("AI Improvement:")
                doc.add_paragraph(data['optimized_content'][:500])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except ImportError:
        return None
    except Exception as e:
        import traceback; traceback.print_exc()
        return None


def export_to_pdf(results: Dict[str, Any]) -> Optional[bytes]:
    """Export as PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle('Title', parent=styles['Title'],
                                     textColor=colors.HexColor('#1565C0'), fontSize=20)
        h1_style = ParagraphStyle('H1', parent=styles['Heading1'],
                                  textColor=colors.HexColor('#1565C0'), fontSize=14)
        body = styles['BodyText']

        summary = results.get("summary", {})
        score = summary.get("overall_score", 0)

        story.append(Paragraph("Karoo v2.0 â€” Optimization Report", title_style))
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')} | "
                                f"AI: {results.get('llm_provider','Rule-Based')}", body))
        story.append(Spacer(1, 0.5*cm))

        # Score table
        score_data = [
            ["Metric", "Value"],
            ["Overall ATS Score", f"{score}/100"],
            ["Interview Probability", f"{summary.get('interview_probability',0)}%"],
            ["Recommended Variant", summary.get('recommended_variant','BALANCED')],
            ["Verdict", summary.get('verdict','')],
        ]
        t = Table(score_data, colWidths=[8*cm, 10*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1565C0')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0f4ff')]),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

        # Agent scores
        story.append(Paragraph("Agent Scores", h1_style))
        agent_rows = [["Agent", "Score", "Rating"]]
        for name, s in summary.get('agent_scores', {}).items():
            rating = "Excellent" if s>=85 else "Strong" if s>=75 else "Adequate" if s>=60 else "Needs Work"
            agent_rows.append([name.replace('_',' ').title(), f"{s}/100", rating])
        at = Table(agent_rows, colWidths=[9*cm, 4*cm, 5*cm])
        at.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1565C0')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0f4ff')]),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('PADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(at)
        story.append(Spacer(1, 0.5*cm))

        # Action items
        story.append(Paragraph("Priority Action Items", h1_style))
        for i, item in enumerate(results.get('action_items', [])[:15], 1):
            story.append(Paragraph(f"{i}. {item}", body))
        story.append(Spacer(1, 0.5*cm))

        # Cover letter
        cl = results.get('cover_letter', '')
        if cl:
            story.append(Paragraph("Cover Letter", h1_style))
            for para in cl.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip().replace('\n', '<br/>'), body))
                    story.append(Spacer(1, 0.2*cm))

        doc.build(story)
        buf.seek(0)
        return buf.getvalue()

    except ImportError:
        return None
    except Exception as e:
        return None
