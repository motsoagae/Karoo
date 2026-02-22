"""
Karoo v2.0 — Main Streamlit Application
11 AI Agents | Full CV Rewrite | Interview Coach | Salary Intel | Job URL Scraper
Groq (free) + OpenAI + Anthropic
"""
import asyncio
import hashlib
import logging
import os
import sys
from io import BytesIO
from datetime import datetime

import streamlit as st
from dotenv import load_dotenv

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Karoo v2 | AI CV Optimizer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ──────────────────────────────────────────────────────────────────────
st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;900&display=swap');
* { font-family: 'Inter', sans-serif; }

.main-header {
    font-size: 3.2rem; font-weight: 900; text-align: center;
    background: linear-gradient(135deg, #1565C0, #0D47A1, #1976D2);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    margin-bottom: 0; line-height: 1.1;
}
.version-badge {
    text-align: center; margin-bottom: 0.5rem;
}
.version-badge span {
    background: #1565C0; color: white; padding: 2px 12px;
    border-radius: 20px; font-size: 0.75rem; font-weight: 700;
    letter-spacing: 1px;
}
.sub-header {
    font-size: 1rem; color: #666; text-align: center; margin-bottom: 1.5rem;
}
.score-card {
    background: linear-gradient(135deg, #1565C0, #0D47A1);
    color: white; border-radius: 16px; padding: 24px; text-align: center;
    margin-bottom: 10px; box-shadow: 0 4px 20px rgba(21,101,192,0.3);
}
.score-number { font-size: 4rem; font-weight: 900; line-height: 1; }
.score-label { font-size: 0.8rem; opacity: 0.8; margin-bottom: 4px; }

.agent-card {
    border-left: 4px solid #1565C0; padding: 10px 14px;
    margin: 6px 0; background: #f7f9ff; border-radius: 0 10px 10px 0;
    transition: all 0.2s;
}
.agent-card:hover { background: #eef2ff; box-shadow: 2px 2px 8px rgba(0,0,0,0.06); }

.action-item-critical {
    background: #ffebee; border-left: 4px solid #C62828;
    padding: 8px 12px; border-radius: 0 8px 8px 0; margin: 4px 0;
    font-size: 0.9rem;
}
.action-item-important {
    background: #fff3e0; border-left: 4px solid #E65C00;
    padding: 8px 12px; border-radius: 0 8px 8px 0; margin: 4px 0;
    font-size: 0.9rem;
}
.action-item-helpful {
    background: #e8f5e9; border-left: 4px solid #2E7D32;
    padding: 8px 12px; border-radius: 0 8px 8px 0; margin: 4px 0;
    font-size: 0.9rem;
}

.feature-pill {
    display: inline-block; background: #e3f2fd; color: #1565C0;
    padding: 3px 10px; border-radius: 20px; font-size: 0.75rem;
    font-weight: 600; margin: 2px; border: 1px solid #90caf9;
}
.new-badge {
    background: #ff6f00; color: white; padding: 1px 6px;
    border-radius: 10px; font-size: 0.65rem; font-weight: 700;
    vertical-align: top; margin-left: 4px;
}

.verdict-box {
    background: linear-gradient(135deg, #e3f2fd, #bbdefb);
    border-left: 5px solid #1565C0; padding: 14px 18px;
    border-radius: 0 12px 12px 0; margin: 10px 0;
    font-size: 1rem; font-weight: 600; color: #1565C0;
}
.stButton>button {
    font-size: 1.05rem; font-weight: 700; padding: 14px;
    transition: all 0.2s;
}
.stButton>button:hover { transform: translateY(-1px); }
.tab-content { padding: 10px 0; }

/* Progress bar custom */
.stProgress > div > div > div { background-color: #1565C0; }
</style>""", unsafe_allow_html=True)


# ─── File Extraction ───────────────────────────────────────────────────────────

def read_pdf(file) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        if len(text.strip()) >= 100:
            return text
    except Exception:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(file) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    except Exception as e:
        st.error(f"PDF read failed: {e}. Try pasting text instead.")
        return ""


def read_docx(file) -> str:
    try:
        from docx import Document
        doc = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        return "\n".join(parts)
    except Exception as e:
        st.error(f"DOCX read failed: {e}")
        return ""


# ─── API Detection ─────────────────────────────────────────────────────────────

def detect_available_llm():
    checks = [
        ("GROQ_API_KEY", "gsk_your", "Groq (free)"),
        ("OPENAI_API_KEY", "sk-your", "OpenAI"),
        ("ANTHROPIC_API_KEY", "sk-ant-your", "Anthropic"),
    ]
    for env_var, placeholder, label in checks:
        val = os.getenv(env_var, "")
        if val and not val.startswith(placeholder):
            return True, label
    return False, None


# ─── Sidebar ───────────────────────────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div style='text-align:center;padding:10px 0 5px'>
            <span style='font-size:2rem'>🤖</span>
            <div style='font-weight:900;font-size:1.1rem;color:#1565C0'>Karoo v2.0</div>
        </div>
        """, unsafe_allow_html=True)

        has_key, provider = detect_available_llm()
        if has_key:
            st.success(f"🟢 {provider} Connected")
        else:
            st.error("🔴 No API Key")
            with st.expander("How to get free AI (30 sec)"):
                st.markdown("""
                **Groq — completely free:**
                1. Visit [console.groq.com](https://console.groq.com)
                2. Sign up (no credit card)
                3. Create API key
                4. Add to `.env` as `GROQ_API_KEY=your_key`
                5. Restart the app
                """)

        st.divider()
        st.subheader("⚙️ Settings")

        target_market = st.selectbox(
            "🌍 Target Market",
            ["Both", "South Africa", "International"],
            help="Adjusts agent weights and market-specific optimization"
        )
        experience_level = st.selectbox(
            "📈 Experience Level",
            ["Entry (0-2 yrs)", "Mid (3-7 yrs)", "Senior (8-15 yrs)", "Executive (15+ yrs)"],
            index=1
        )
        industry = st.text_input("🏭 Industry", placeholder="e.g. FinTech, Mining, Healthcare")
        target_role = st.text_input("🎯 Target Role", placeholder="e.g. Senior Data Analyst")

        st.divider()
        st.subheader("🛠️ Features")

        generate_cl = st.checkbox("✉️ Generate Cover Letter", value=True)
        rewrite_cv = st.checkbox(
            "✍️ AI Full CV Rewrite",
            value=True,
            help="Uses AI to completely rewrite your CV in all 3 variants. Requires API key."
        )
        run_interview = st.checkbox(
            "🎤 Interview Coach",
            value=True,
            help="Generates likely interview questions + STAR frameworks"
        )
        run_salary = st.checkbox(
            "💰 Salary Intelligence",
            value=True,
            help="Market rate analysis + negotiation strategy"
        )

        st.divider()
        st.caption("**v2 New Agents:**")
        st.caption("🎤 Interview Coach | 💰 Salary Intel")
        st.caption("✍️ AI CV Rewriter | 🔗 Job URL Scraper")
        st.caption("📊 Visual Analytics | 📄 PDF Export")

        return {
            "target_market": target_market,
            "experience_level": experience_level.split(" ")[0],
            "industry": industry,
            "target_role": target_role,
            "generate_cover_letter": generate_cl,
            "rewrite_cv": rewrite_cv,
            "run_interview": run_interview,
            "run_salary": run_salary,
        }


# ─── Results ───────────────────────────────────────────────────────────────────

def render_results(results: dict):
    summary = results.get("summary", {})
    overall = summary.get("overall_score", 0)
    prob = summary.get("interview_probability", 0)
    agent_scores = summary.get("agent_scores", {})

    # ── Hero Dashboard ──────────────────────────────────────────────
    st.subheader("📊 Performance Dashboard")
    col_gauge, col_metrics = st.columns([1, 2])

    with col_gauge:
        try:
            from src.ui.charts import score_gauge
            fig = score_gauge(int(overall))
            if fig:
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.markdown(f"""<div class="score-card">
                    <div class="score-label">ATS SCORE</div>
                    <div class="score-number">{int(overall)}</div>
                    <div class="score-label">/100</div>
                </div>""", unsafe_allow_html=True)
        except Exception:
            st.markdown(f"""<div class="score-card">
                <div class="score-label">ATS SCORE</div>
                <div class="score-number">{int(overall)}</div>
                <div class="score-label">/100</div>
            </div>""", unsafe_allow_html=True)

    with col_metrics:
        m1, m2 = st.columns(2)
        delta_prob = f"+{prob-50}% vs avg" if prob > 50 else f"{prob-50}% vs avg"
        m1.metric("🎯 Interview Probability", f"{prob}%", delta_prob)
        m1.metric("⭐ Recommended Variant", summary.get("recommended_variant", "BALANCED"))
        m2.metric("🤖 AI Engine", results.get("llm_provider", "Rule-Based"))
        m2.metric("⏱️ Analysis Time", f"{results.get('metadata',{}).get('execution_seconds',0)}s")

        model = results.get("llm_model", "N/A")
        ai_count = summary.get("ai_powered_count", 0)
        st.caption(f"Model: `{model}` · AI-Powered: {ai_count}/10 agents")

    st.markdown(f'<div class="verdict-box">📋 {summary.get("verdict","")}</div>', unsafe_allow_html=True)

    # ── Visualisations ──────────────────────────────────────────────
    if agent_scores:
        st.subheader("📈 Score Analytics")
        vcol1, vcol2 = st.columns(2)
        try:
            from src.ui.charts import radar_chart, agent_bar_chart
            with vcol1:
                fig = radar_chart(agent_scores)
                if fig: st.plotly_chart(fig, use_container_width=True)
            with vcol2:
                fig = agent_bar_chart(agent_scores)
                if fig: st.plotly_chart(fig, use_container_width=True)
        except Exception:
            pass

    # ── Agent Grid ──────────────────────────────────────────────────
    st.subheader("🤖 Agent Scores")
    ICONS = {
        "algorithm_breaker": "🎯", "sa_specialist": "🇿🇦", "global_setter": "🌍",
        "recruiter_scanner": "👁️", "hiring_manager": "💼", "semantic_matcher": "📊",
        "compliance_guardian": "⚖️", "future_architect": "🚀",
        "interview_coach": "🎤", "salary_intelligence": "💰",
    }
    cols = st.columns(5)
    for i, (name, s) in enumerate(agent_scores.items()):
        with cols[i % 5]:
            icon = ICONS.get(name, "🤖")
            label = name.replace("_", " ").title()
            color = "#2E7D32" if s >= 80 else "#E65C00" if s >= 60 else "#C62828"
            bar = "█" * (s // 10) + "░" * (10 - s // 10)
            st.markdown(f"""<div class="agent-card">
                <b>{icon} {label}</b><br>
                <span style="color:{color};font-size:1.4rem;font-weight:900">{s}</span>
                <span style="color:#888;font-size:0.8rem">/100</span><br>
                <span style="font-size:0.72rem;color:#aaa;font-family:monospace">{bar}</span>
            </div>""", unsafe_allow_html=True)

    # ── Action Items ────────────────────────────────────────────────
    st.subheader("✅ Priority Action Items")
    items = results.get("action_items", [])
    if items:
        for i, item in enumerate(items, 1):
            if i <= 4:
                st.markdown(f'<div class="action-item-critical">🔴 <b>{i}.</b> {item}</div>', unsafe_allow_html=True)
            elif i <= 9:
                st.markdown(f'<div class="action-item-important">🟡 <b>{i}.</b> {item}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="action-item-helpful">🟢 <b>{i}.</b> {item}</div>', unsafe_allow_html=True)
    else:
        st.info("No action items — add an API key for AI-powered analysis.")

    # ── CV Variants ────────────────────────────────────────────────
    st.subheader("📄 Optimized CV Variants")
    variants = results.get("cv_variants", {})
    ai_rewrite = results.get("metadata", {}).get("ai_rewrite", False)

    if ai_rewrite:
        st.success("✨ **AI-Rewritten CVs:** Full rewrites generated by AI — review all facts before sending.")
    else:
        st.info("💡 **Enhancement Mode:** CV variants show optimized sections. Add API key for full AI rewrites.")

    t1, t2, t3 = st.tabs(["🎯 ATS-MAX", "⭐ BALANCED (Recommended)", "🎨 CREATIVE"])
    with t1:
        st.info("Maximum ATS parse rate. Best for large corporations, strict ATS portals.")
        content = variants.get("ats_max", "")
        st.text_area("ATS-MAX CV", content, height=500, key="v_ats")
        if content:
            st.download_button("📋 Copy ATS-MAX", content, "cv_ats_max.txt", use_container_width=True)
    with t2:
        st.success("Best of both worlds — ATS-optimized + human-readable. For 85% of applications.")
        content = variants.get("balanced", "")
        st.text_area("BALANCED CV", content, height=500, key="v_bal")
        if content:
            st.download_button("📋 Copy Balanced", content, "cv_balanced.txt", use_container_width=True)
    with t3:
        st.info("Human-first. Best for startups, agencies, executive search, creative roles.")
        content = variants.get("creative", "")
        st.text_area("CREATIVE CV", content, height=500, key="v_cre")
        if content:
            st.download_button("📋 Copy Creative", content, "cv_creative.txt", use_container_width=True)

    # ── Cover Letter ───────────────────────────────────────────────
    cl = results.get("cover_letter", "")
    if cl:
        st.subheader("✉️ Cover Letter")
        word_count = len(cl.split())
        col_cl1, col_cl2 = st.columns([3, 1])
        with col_cl1:
            st.caption(f"📊 {word_count} words · Customize company name and date before sending")
        with col_cl2:
            st.download_button("📋 Copy Letter", cl, "cover_letter.txt", use_container_width=True)
        st.text_area("Cover Letter", cl, height=350, key="cover_letter_area")

    # ── Interview Coach Tab ────────────────────────────────────────
    interview_data = results.get("agent_results", {}).get("interview_coach", {})
    if interview_data.get("optimized_content"):
        st.subheader("🎤 Interview Preparation")
        with st.expander("View Interview Questions + STAR Frameworks", expanded=False):
            st.write(f"**Readiness Score:** {interview_data.get('score', 0)}/100")
            st.code(interview_data["optimized_content"], language=None)

    # ── Salary Intelligence ────────────────────────────────────────
    salary_data = results.get("agent_results", {}).get("salary_intelligence", {})
    if salary_data.get("findings"):
        st.subheader("💰 Salary Intelligence")
        salary_cols = st.columns(2)
        with salary_cols[0]:
            for f in salary_data.get("findings", [])[:4]:
                st.write(f"• {f}")
        with salary_cols[1]:
            if salary_data.get("optimized_content"):
                st.info(f"💬 **Negotiation Script:**\n\n{salary_data['optimized_content'][:400]}")

    # ── Full Agent Reports ─────────────────────────────────────────
    with st.expander("🔍 Full Agent Reports (all 10 agents)", expanded=False):
        for name, data in results.get("agent_results", {}).items():
            icon = ICONS.get(name, "🤖")
            s = data.get("score", 0)
            ai_tag = "🧠 AI" if data.get("ai_powered") else "📐 Rule-Based"
            with st.expander(f"{icon} {name.replace('_',' ').title()} — {s}/100 ({ai_tag})"):
                st.write("**Findings:**")
                for f in data.get("findings", []):
                    st.write(f"• {f}")
                st.write("**Recommendations:**")
                for r in data.get("recommendations", []):
                    st.write(f"→ {r}")
                if data.get("optimized_content") and name not in ["interview_coach"]:
                    st.success(f"💡 **AI Enhancement:** {data['optimized_content'][:400]}")

    # ── Downloads ──────────────────────────────────────────────────
    st.subheader("💾 Export Results")
    d1, d2, d3 = st.columns(3)
    ts = datetime.now().strftime('%Y%m%d_%H%M')

    with d1:
        from src.core.exporter import export_to_txt
        txt = export_to_txt(results)
        st.download_button(
            "📄 Full Report (TXT)",
            data=txt,
            file_name=f"Karoo_v2_{ts}.txt",
            mime="text/plain",
            use_container_width=True
        )

    with d2:
        try:
            from src.core.exporter import export_to_docx
            docx = export_to_docx(results)
            if docx:
                st.download_button(
                    "📝 Full Report (DOCX)",
                    data=docx,
                    file_name=f"Karoo_v2_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.info("DOCX: `pip install python-docx`")
        except Exception:
            st.info("DOCX export unavailable")

    with d3:
        try:
            from src.core.exporter import export_to_pdf
            pdf = export_to_pdf(results)
            if pdf:
                st.download_button(
                    "🔴 Full Report (PDF)",
                    data=pdf,
                    file_name=f"ats_god_v2_{ts}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            else:
                st.info("PDF: `pip install reportlab`")
        except Exception:
            st.info("PDF export unavailable")


# ─── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Header
    st.markdown('<h1 class="main-header">🤖 Karoo</h1>', unsafe_allow_html=True)
    st.markdown('<div class="version-badge"><span>VERSION 2.0</span></div>', unsafe_allow_html=True)
    st.markdown("""<p class="sub-header">
        11 AI Agents · Full CV Rewrite · Interview Coach · Salary Intelligence · Job URL Scraper<br>
        <span class="feature-pill">South Africa</span>
        <span class="feature-pill">International</span>
        <span class="feature-pill">Groq Free</span>
        <span class="feature-pill">OpenAI</span>
        <span class="feature-pill">Anthropic</span>
        <span class="feature-pill">🆕 AI Rewrite</span>
        <span class="feature-pill">🆕 Interview Coach</span>
        <span class="feature-pill">🆕 Salary Intel</span>
    </p>""", unsafe_allow_html=True)

    context = render_sidebar()

    # ── Input Section ──────────────────────────────────────────────
    col_cv, col_jd = st.columns(2)

    with col_cv:
        st.subheader("📤 Your CV")
        method = st.radio("Input method", ["Upload PDF/DOCX", "Paste text"], horizontal=True, key="cv_method")
        cv_text = ""

        if method == "Upload PDF/DOCX":
            uploaded = st.file_uploader("Upload CV", type=["pdf", "docx"], key="cv_upload")
            if uploaded:
                size_mb = len(uploaded.getvalue()) / (1024 * 1024)
                if size_mb > 10:
                    st.error("Max 10MB")
                else:
                    buf = BytesIO(uploaded.read())
                    with st.spinner("Extracting text..."):
                        cv_text = read_pdf(buf) if "pdf" in uploaded.type else read_docx(buf)
                    if cv_text:
                        wc = len(cv_text.split())
                        st.success(f"✓ {wc:,} words · {size_mb:.1f}MB · {uploaded.name}")
                        if wc < 100:
                            st.warning("Very short CV — scanned image? Try pasting text.")
                        with st.expander("Preview (first 1000 chars)"):
                            st.text(cv_text[:1000] + ("…" if len(cv_text) > 1000 else ""))
        else:
            cv_text = st.text_area(
                "Paste CV text",
                height=400,
                placeholder="Paste your complete CV here…",
                key="cv_paste"
            )

    with col_jd:
        st.subheader("📋 Job Description")

        jd_input_method = st.radio(
            "JD input method",
            ["Paste text", "🔗 Paste URL (auto-extract)"],
            horizontal=True, key="jd_method"
        )
        jd_text = ""

        if "URL" in jd_input_method:
            url = st.text_input(
                "Job posting URL",
                placeholder="https://www.linkedin.com/jobs/... or https://pnet.co.za/...",
                key="jd_url"
            )
            if url:
                from src.core.job_scraper import scrape_job_url, is_valid_url
                if is_valid_url(url):
                    if st.button("🔗 Extract Job Description", key="scrape_btn"):
                        with st.spinner("Scraping job posting..."):
                            scraped, platform = scrape_job_url(url)
                        if scraped:
                            st.session_state["scraped_jd"] = scraped
                            st.session_state["scraped_platform"] = platform
                            st.success(f"✓ Extracted from {platform} ({len(scraped.split()):,} words)")
                        else:
                            st.error(f"Could not extract JD. Paste manually instead. ({platform})")
                else:
                    st.warning("Please enter a valid URL starting with https://")

                if st.session_state.get("scraped_jd"):
                    jd_text = st.text_area(
                        "Extracted JD (edit if needed)",
                        value=st.session_state["scraped_jd"],
                        height=380, key="jd_scraped_edit"
                    )
        else:
            jd_text = st.text_area(
                "Paste job description",
                height=420,
                placeholder="Paste the complete job description here…",
                key="jd_paste"
            )

    st.divider()

    # ── Validation ─────────────────────────────────────────────────
    cv_ok = bool(cv_text and cv_text.strip())
    jd_ok = bool(jd_text and jd_text.strip())

    if not cv_ok:
        st.warning("👆 Add your CV — upload a file or paste text above")
    if not jd_ok:
        st.warning("👆 Add the job description — paste or use URL extraction")

    if cv_ok and jd_ok:
        wc_cv = len(cv_text.split())
        wc_jd = len(jd_text.split())

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("CV Words", f"{wc_cv:,}",
                  "✓ Good length" if 250 <= wc_cv <= 1200 else ("Too short" if wc_cv < 250 else "May be too long"))
        m2.metric("JD Words", f"{wc_jd:,}",
                  "✓ Detailed" if wc_jd >= 100 else "More detail = better analysis")
        has_key, provider = detect_available_llm()
        m3.metric("AI Mode", provider or "Rule-Based",
                  "Full AI — 11 agents" if has_key else "Add GROQ_API_KEY for AI")
        ai_features = sum([
            context.get("generate_cover_letter", True),
            context.get("rewrite_cv", True),
            context.get("run_interview", True),
            context.get("run_salary", True),
        ])
        m4.metric("Features Active", f"{ai_features}/4",
                  "All features on" if ai_features == 4 else f"{4-ai_features} features off")

        # Launch button
        agent_count = 8 + context.get("run_interview", True) + context.get("run_salary", True)
        btn_label = f"🚀 OPTIMIZE — LAUNCH {agent_count} AI AGENTS"
        if st.button(btn_label, type="primary", use_container_width=True):

            prog = st.progress(0.0)
            status = st.empty()

            def cb(pct: float, msg: str):
                prog.progress(min(pct, 1.0))
                status.text(msg)

            try:
                sys.path.insert(0, ".")
                from src.core.orchestrator import KarooOrchestrator

                # Filter agents based on feature flags
                orch = KarooOrchestrator()
                if not context.get("run_interview", True):
                    orch.agents.pop("interview_coach", None)
                if not context.get("run_salary", True):
                    orch.agents.pop("salary_intelligence", None)

                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)

                results = loop.run_until_complete(
                    orch.optimize(
                        cv_text=cv_text,
                        job_description=jd_text,
                        context=context,
                        generate_cover_letter=context.get("generate_cover_letter", True),
                        rewrite_cv=context.get("rewrite_cv", True),
                        progress_callback=cb,
                    )
                )
                loop.close()

                prog.progress(1.0)
                status.text("✅ All agents complete!")
                st.balloons()

                elapsed = results.get("metadata", {}).get("execution_seconds", 0)
                provider = results.get("llm_provider", "Rule-Based")
                agents_run = results.get("metadata", {}).get("agents_run", 0)
                st.success(f"✅ {agents_run} agents complete in {elapsed}s — {provider} mode")

                st.divider()
                render_results(results)

                # Store in session for re-display
                st.session_state["last_results"] = results

            except Exception as e:
                st.error(f"❌ Optimization failed: {e}")
                import traceback
                with st.expander("Debug info (for developers)"):
                    st.code(traceback.format_exc())
                st.info("Ensure all files are in `src/` and run: `export PYTHONPATH=.`")

    # ── Previous Results ───────────────────────────────────────────
    elif st.session_state.get("last_results") and not (cv_ok and jd_ok):
        st.info("📊 Showing previous analysis results:")
        render_results(st.session_state["last_results"])


if __name__ == "__main__":
    main()
